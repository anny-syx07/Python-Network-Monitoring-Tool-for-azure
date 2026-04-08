#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Script tạo báo cáo đồ án .docx với định dạng chuẩn."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ===== SETUP PAGE: A4, margins =====
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# ===== STYLE SETUP =====
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5

# Set font for East Asian
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
if rFonts is None:
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rPr.append(rFonts)

def set_run_font(run, size=13, bold=False, italic=False, underline=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rPr.insert(0, rFonts)

def add_paragraph(text, size=13, bold=False, italic=False, alignment=None, space_after=None, space_before=None, underline=False):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, underline=underline)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_chapter_title(text):
    """Tiêu đề chương: 16pt, in đậm, căn giữa"""
    p = add_paragraph(text, size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=12)
    return p

def add_section_title(text):
    """Tiêu đề mục: 14pt, in đậm"""
    return add_paragraph(text, size=14, bold=True, space_before=12, space_after=6)

def add_subsection_title(text):
    """Tiêu đề tiểu mục: 14pt, in nghiêng"""
    return add_paragraph(text, size=14, italic=True, space_before=8, space_after=4)

def add_subsubsection_title(text):
    """Tiểu mục con: 13pt, gạch chân"""
    return add_paragraph(text, size=13, underline=True, space_before=6, space_after=4)

def add_body(text, space_after=6):
    return add_paragraph(text, size=13, space_after=space_after)

def add_bullet(text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    run = p.add_run(f"• {text}")
    set_run_font(run, size=13)
    return p

def add_code_block(code_text):
    """Add code as a formatted paragraph"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    return p

def add_table(headers, rows, caption=None):
    if caption:
        p = add_paragraph(caption, size=13, bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=12, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=12)
    doc.add_paragraph()  # spacing
    return table

def add_page_break():
    doc.add_page_break()

def add_figure_caption(text):
    p = add_paragraph(text, size=13, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8)
    return p

# ============================================================
# TRANG BÌA
# ============================================================
add_paragraph("BỘ GIÁO DỤC VÀ ĐÀO TẠO", size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_paragraph("TRƯỜNG ĐẠI HỌC QUẢN LÝ & CÔNG NGHỆ TP. HCM", size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_paragraph("KHOA CÔNG NGHỆ THÔNG TIN", size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

add_paragraph("ĐỒ ÁN MÔN HỌC", size=20, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=48, space_after=24)

add_paragraph("XÂY DỰNG CÔNG CỤ GIÁM SÁT", size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_paragraph("MẠNG BẰNG PYTHON VÀ TRIỂN KHAI", size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_paragraph("TRÊN ĐÁM MÂY AZURE", size=18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)

add_paragraph("Ngành: CÔNG NGHỆ THÔNG TIN", size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_paragraph("Chuyên ngành: MẠNG MÁY TÍNH VÀ ĐIỆN TOÁN ĐÁM MÂY", size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_paragraph("Giảng viên hướng dẫn: Nguyễn Thanh Phong", size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_paragraph("Lớp: 1747", size=13, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_table(
    ["Sinh viên thực hiện", "MSSV"],
    [
        ["1. Ngô Đình Thông", "2403700248"],
        ["2. Nguyễn Nhật Duy", "2403700298"],
        ["3. Bùi Thị Anh Thư", "2403700243"],
    ]
)

add_paragraph("TP. Hồ Chí Minh, 2026", size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=36)

add_page_break()

# ============================================================
# LỜI CAM ĐOAN
# ============================================================
add_chapter_title("LỜI CAM ĐOAN")
add_body('Chúng tôi xin cam đoan đồ án môn học "Xây dựng công cụ giám sát mạng bằng Python và triển khai trên đám mây Azure" là công trình nghiên cứu của nhóm chúng tôi dưới sự hướng dẫn của thầy Nguyễn Thanh Phong. Các số liệu, kết quả trong đồ án là trung thực. Mọi tham khảo đều được trích dẫn đầy đủ.')
add_body("")
add_paragraph("TP. Hồ Chí Minh, ngày ... tháng ... năm 2026", size=13, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
add_paragraph("Nhóm sinh viên thực hiện", size=13, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

add_page_break()

# ============================================================
# DANH MỤC CHỮ VIẾT TẮT
# ============================================================
add_chapter_title("DANH MỤC CHỮ VIẾT TẮT")

abbreviations = [
    ("API", "Application Programming Interface"),
    ("CSV", "Comma-Separated Values"),
    ("DNS", "Domain Name System"),
    ("GUI", "Graphical User Interface"),
    ("HTTP", "HyperText Transfer Protocol"),
    ("HTTPS", "HTTP Secure"),
    ("ICMP", "Internet Control Message Protocol"),
    ("IP", "Internet Protocol"),
    ("LAN", "Local Area Network"),
    ("NMS", "Network Management System"),
    ("REST", "Representational State Transfer"),
    ("RTT", "Round Trip Time"),
    ("SMB", "Server Message Block"),
    ("TCP", "Transmission Control Protocol"),
    ("TLS", "Transport Layer Security"),
    ("UI", "User Interface"),
]
add_table(["Viết tắt", "Nghĩa đầy đủ"], abbreviations)

add_page_break()

# ============================================================
# MỤC LỤC (placeholder)
# ============================================================
add_chapter_title("MỤC LỤC")
toc_items = [
    "LỜI CAM ĐOAN",
    "DANH MỤC CHỮ VIẾT TẮT",
    "DANH MỤC CÁC BẢNG",
    "DANH MỤC CÁC HÌNH VẼ",
    "",
    "Chương 1. MỞ ĐẦU",
    "    1.1. Đặt vấn đề",
    "    1.2. Mục tiêu đề tài",
    "    1.3. Phạm vi và giới hạn",
    "    1.4. Bố cục đồ án",
    "Chương 2. CƠ SỞ LÝ THUYẾT",
    "    2.1. Tổng quan về giám sát mạng",
    "    2.2. Giao thức mạng sử dụng trong đề tài",
    "    2.3. Lập trình mạng với Python",
    "    2.4. Giao diện đồ họa với Tkinter",
    "    2.5. Trực quan hóa dữ liệu với Matplotlib",
    "    2.6. Framework Flask và RESTful API",
    "    2.7. Điện toán đám mây với Microsoft Azure",
    "Chương 3. KẾT QUẢ THỰC NGHIỆM",
    "    3.1. Thiết kế hệ thống",
    "    3.2. Cài đặt chi tiết – Ứng dụng Desktop",
    "    3.3. Cài đặt chi tiết – API Cloud",
    "    3.4. Triển khai lên Azure",
    "    3.5. Kết quả thu được",
    "Chương 4. KẾT LUẬN VÀ KIẾN NGHỊ",
    "    4.1. Kết luận",
    "    4.2. Kiến nghị và hướng phát triển",
    "",
    "TÀI LIỆU THAM KHẢO",
    "PHỤ LỤC",
]
for item in toc_items:
    if item == "":
        add_body("")
    elif item.startswith("    "):
        add_body(item, space_after=2)
    else:
        p = add_paragraph(item, size=13, bold=True, space_after=2)

add_page_break()

# ============================================================
# DANH MỤC CÁC BẢNG
# ============================================================
add_chapter_title("DANH MỤC CÁC BẢNG")
tables_list = [
    "Bảng 1.1. Phạm vi đề tài",
    "Bảng 2.1. So sánh các công cụ giám sát mạng phổ biến",
    "Bảng 3.1. Mô tả các file trong dự án",
    "Bảng 3.2. Các thành phần giao diện GUI",
    "Bảng 3.3. Cấu hình tài nguyên Azure đã triển khai",
    "Bảng 3.4. Thống kê dữ liệu thực nghiệm",
    "Bảng 3.5. Tổng kết đánh giá các chỉ tiêu",
    "Bảng 4.1. Hạn chế hiện tại",
]
for t in tables_list:
    add_body(t, space_after=2)

add_chapter_title("DANH MỤC CÁC HÌNH VẼ")
figures_list = [
    "Hình 2.1. Mô hình hoạt động giao thức ICMP",
    "Hình 3.1. Kiến trúc tổng quan hệ thống giám sát mạng",
    "Hình 3.2. Giao diện ứng dụng Network Monitor PRO",
    "Hình 3.3. Biểu đồ giám sát mạng thời gian thực",
    "Hình 3.4. Kết quả quét nhanh mạng LAN",
    "Hình 3.5. Dữ liệu CSV trả về qua trình duyệt",
    "Hình 3.6. File Share csvreports trên Azure Portal",
]
for f in figures_list:
    add_body(f, space_after=2)

add_page_break()

# ============================================================
# CHƯƠNG 1. MỞ ĐẦU
# ============================================================
add_chapter_title("Chương 1. MỞ ĐẦU")

add_section_title("1.1. Đặt vấn đề")
add_body('Trong bối cảnh hạ tầng công nghệ thông tin ngày càng phức tạp và phân tán, việc giám sát liên tục trạng thái hoạt động của các thiết bị và dịch vụ mạng trở thành yêu cầu cấp thiết đối với mọi tổ chức. Sự cố mạng nếu không được phát hiện và xử lý kịp thời có thể gây ra gián đoạn dịch vụ, tổn thất kinh tế và mất uy tín đối với người dùng [1].')
add_body('Các hệ thống quản lý mạng chuyên nghiệp hiện có như Nagios, Zabbix hay PRTG Network Monitor cung cấp khả năng giám sát toàn diện, nhưng đi kèm với chi phí triển khai cao, yêu cầu máy chủ chuyên dụng và độ phức tạp trong cấu hình. Điều này tạo ra rào cản đáng kể cho các tổ chức vừa và nhỏ hoặc các nhóm kỹ thuật có nguồn lực hạn chế.')
add_body('Python là ngôn ngữ lập trình mạnh mẽ, phổ biến và sở hữu hệ sinh thái thư viện phong phú, đặc biệt là các thư viện hỗ trợ lập trình mạng (ping3, requests), giao diện đồ họa (Tkinter) và trực quan hóa dữ liệu (Matplotlib) [2]. Kết hợp với nền tảng đám mây Microsoft Azure, hoàn toàn có thể xây dựng một công cụ giám sát mạng có giao diện trực quan, khả năng lưu trữ đám mây và hoạt động liên tục mà không cần đầu tư phần cứng.')
add_body('Từ những phân tích trên, đề tài "Xây dựng công cụ giám sát mạng bằng Python và triển khai trên nền tảng đám mây Microsoft Azure" được đề xuất nhằm giải quyết bài toán giám sát mạng với chi phí thấp, giao diện thân thiện và khả năng lưu trữ dữ liệu tập trung trên Cloud.')

add_section_title("1.2. Mục tiêu đề tài")
add_body("Đề tài hướng đến các mục tiêu cụ thể sau:")
add_subsection_title("Mục tiêu chính:")
add_bullet("Xây dựng ứng dụng giám sát mạng bằng Python có giao diện đồ họa (GUI), có khả năng ping kiểm tra trạng thái host theo thời gian thực, hiển thị biểu đồ trực quan và quét mạng LAN.")
add_subsection_title("Mục tiêu phụ:")
add_bullet("Xây dựng API trên nền tảng Flask để tiếp nhận và lưu trữ dữ liệu giám sát tập trung trên đám mây Azure.")
add_bullet("Áp dụng kỹ thuật lập trình đa luồng (multithreading) để ứng dụng không bị đóng băng khi ping.")
add_bullet("Triển khai biểu đồ thời gian thực (real-time chart) bằng Matplotlib Animation.")
add_bullet("Bảo mật API bằng cơ chế xác thực Bearer Token.")
add_bullet("Lưu trữ bền vững dữ liệu CSV trên Azure File Share thông qua Azure App Service.")

add_section_title("1.3. Phạm vi và giới hạn")
add_body("Phạm vi thực hiện:")
add_table(
    ["Khía cạnh", "Chi tiết"],
    [
        ["Ngôn ngữ lập trình", "Python 3 (Tkinter, Flask, ping3, Matplotlib)"],
        ["Nền tảng đám mây", "Microsoft Azure (Azure for Students)"],
        ["Dịch vụ Azure", "App Service (B1), Storage Account, File Share"],
        ["Giao thức giám sát", "ICMP (ping)"],
        ["Mục tiêu giám sát", "Địa chỉ IP tùy chọn (mặc định 8.8.8.8) + quét LAN"],
        ["Đầu ra", "File CSV (local + cloud), biểu đồ thời gian thực"],
        ["Tổng số bản ghi", "14.765+ bản ghi log"],
    ],
    caption="Bảng 1.1. Phạm vi đề tài"
)

add_body("Giới hạn:")
add_bullet("Ứng dụng Desktop yêu cầu chạy trên máy tính có cài Python (chưa đóng gói thành file .exe/.app).")
add_bullet("Phạm vi giám sát giới hạn ở tầng mạng (Layer 3); chưa hỗ trợ giao thức ứng dụng (Layer 7).")
add_bullet("Quét LAN giới hạn ở dải IP .1 đến .10 để đảm bảo tốc độ nhanh.")

add_section_title("1.4. Bố cục đồ án")
add_body("Ngoài phần mở đầu, tài liệu tham khảo và phụ lục, đồ án được tổ chức thành 4 chương:")
add_bullet("Chương 1 – Mở đầu: Trình bày bối cảnh, lý do chọn đề tài, mục tiêu và phạm vi nghiên cứu.")
add_bullet("Chương 2 – Cơ sở lý thuyết: Tổng quan các giao thức mạng, kỹ thuật lập trình Python, Tkinter, Matplotlib, Flask và nền tảng Azure.")
add_bullet("Chương 3 – Kết quả thực nghiệm: Trình bày thiết kế hệ thống, chi tiết cài đặt, quy trình triển khai lên Azure và kết quả thu được.")
add_bullet("Chương 4 – Kết luận và kiến nghị: Đánh giá kết quả, hạn chế và hướng phát triển tiếp theo.")

add_page_break()

# ============================================================
# CHƯƠNG 2. CƠ SỞ LÝ THUYẾT
# ============================================================
add_chapter_title("Chương 2. CƠ SỞ LÝ THUYẾT")

add_section_title("2.1. Tổng quan về giám sát mạng")
add_subsection_title("2.1.1. Khái niệm giám sát mạng")
add_body("Giám sát mạng (Network Monitoring) là quá trình theo dõi liên tục trạng thái hoạt động của các thành phần hạ tầng mạng bao gồm: thiết bị đầu cuối (host), thiết bị mạng (router, switch), và các dịch vụ mạng (web server, database) [1]. Mục tiêu chính là phát hiện sớm các bất thường, đo lường hiệu suất và đảm bảo tính khả dụng (availability) của hệ thống.")
add_body("Các chỉ số quan trọng trong giám sát mạng bao gồm:")
add_bullet("Uptime (%): Tỷ lệ thời gian host/dịch vụ hoạt động bình thường.")
add_bullet("Latency (ms): Thời gian trễ từ điểm gửi đến điểm nhận, đo bằng millisecond.")
add_bullet("Packet Loss (%): Tỷ lệ gói tin bị mất trên đường truyền.")
add_bullet("Downtime: Khoảng thời gian host/dịch vụ ngừng hoạt động.")

add_subsection_title("2.1.2. Các công cụ giám sát mạng hiện có")
add_table(
    ["Công cụ", "Ưu điểm", "Nhược điểm"],
    [
        ["Nagios", "Mạnh mẽ, plugin phong phú", "Phức tạp, tốn tài nguyên"],
        ["Zabbix", "Dashboard trực quan", "Yêu cầu cấu hình cao"],
        ["PRTG", "Giao diện thân thiện", "Chi phí bản quyền cao"],
        ["Network Monitor PRO\n(đề tài này)", "Nhẹ, miễn phí, có GUI,\ntích hợp Cloud", "Chưa hỗ trợ Layer 7"],
    ],
    caption="Bảng 2.1. So sánh các công cụ giám sát mạng phổ biến"
)

add_section_title("2.2. Giao thức mạng sử dụng trong đề tài")
add_subsection_title("2.2.1. Giao thức ICMP")
add_body("ICMP (Internet Control Message Protocol) là giao thức tầng Network (Layer 3) trong mô hình OSI, được định nghĩa trong RFC 792 [3]. ICMP được thiết kế để truyền thông báo điều khiển và kiểm lỗi giữa các thiết bị mạng.")
add_body("Lệnh ping hoạt động bằng cách:")
add_bullet("Gửi gói tin ICMP Echo Request đến host đích.")
add_bullet("Chờ nhận ICMP Echo Reply từ host đó.")
add_bullet("Tính toán thời gian trễ (RTT – Round Trip Time).")
add_body("Trong đề tài, thư viện ping3 được sử dụng để gửi ICMP ping trực tiếp từ Python, thay vì gọi lệnh hệ điều hành qua subprocess. Điều này giúp code gọn hơn và hoạt động nhất quán trên mọi hệ điều hành.")
add_figure_caption("Hình 2.1. Mô hình hoạt động giao thức ICMP")

add_subsection_title("2.2.2. Giao thức TCP và HTTP/HTTPS")
add_body("TCP (Transmission Control Protocol) là giao thức tầng Transport (Layer 4), cung cấp kết nối tin cậy thông qua cơ chế 3-way handshake [4].")
add_body("Đề tài sử dụng giao thức HTTP/HTTPS (dựa trên TCP) để giao tiếp giữa ứng dụng Desktop và API trên Azure Cloud:")
add_bullet("POST /api/log: Gửi dữ liệu giám sát lên Cloud dưới dạng JSON.")
add_bullet("GET /api/csv: Truy xuất dữ liệu log từ Cloud dưới dạng CSV.")

add_subsection_title("2.2.3. Hệ thống phân giải tên miền DNS")
add_body("DNS (Domain Name System) là hệ thống phân cấp phân giải hostname (google.com) thành địa chỉ IP (172.217.x.x) [5]. Đề tài cho phép người dùng nhập cả hostname hoặc địa chỉ IP vào trường Target IP để kiểm tra.")

add_section_title("2.3. Lập trình mạng với Python")
add_subsection_title("2.3.1. Thư viện ping3")
add_body("ping3 là thư viện Python cho phép gửi ICMP Echo Request và nhận Echo Reply trực tiếp từ Python mà không cần gọi lệnh hệ điều hành [2]. Hàm ping(ip, timeout) trả về thời gian phản hồi (RTT) nếu host online, hoặc None nếu host offline/timeout.")

add_subsection_title("2.3.2. Thư viện requests")
add_body("requests là thư viện HTTP phổ biến nhất trong Python, cung cấp giao diện đơn giản để gửi HTTP request [2]. Đề tài dùng requests.post() để gửi dữ liệu JSON lên Azure API với Bearer Token trong header Authorization.")

add_subsection_title("2.3.3. Module threading và đa luồng")
add_body("Python cung cấp module threading để lập trình đa luồng. Đề tài sử dụng threading.Thread(daemon=True) để tạo luồng nền cho việc ping liên tục và quét LAN, đảm bảo giao diện GUI không bị đóng băng (freeze) trong quá trình xử lý mạng.")

add_subsection_title("2.3.4. Module csv")
add_body("Module csv trong thư viện chuẩn Python hỗ trợ đọc và ghi file CSV (Comma-Separated Values). Đề tài ghi log dự phòng ra file log.csv local bằng csv.writer(), đảm bảo dữ liệu không bị mất nếu kết nối Cloud gặp sự cố.")

add_section_title("2.4. Giao diện đồ họa với Tkinter")
add_body("Tkinter là thư viện GUI tiêu chuẩn đi kèm Python, không cần cài đặt thêm [2]. Tkinter cung cấp các widget cơ bản như Label, Entry, Button, Text, Frame để xây dựng giao diện ứng dụng desktop.")
add_body("Đặc điểm quan trọng của Tkinter liên quan đến đề tài:")
add_bullet("Vòng lặp sự kiện (Event Loop): root.mainloop() chạy vòng lặp chính, lắng nghe sự kiện người dùng.")
add_bullet("Thread-safe update: Tkinter không cho phép cập nhật UI từ luồng phụ. Phải dùng root.after(0, callback) để đẩy việc cập nhật về luồng chính.")
add_bullet("Widget state: Nút bấm có thể vô hiệu hóa (state=tk.DISABLED) để chống thao tác trùng lặp.")

add_section_title("2.5. Trực quan hóa dữ liệu với Matplotlib")
add_body("Matplotlib là thư viện vẽ biểu đồ mạnh mẽ nhất trong Python [6]. Đề tài sử dụng matplotlib.animation.FuncAnimation để tạo biểu đồ cập nhật theo thời gian thực (real-time chart), hiển thị trạng thái Online/Offline của host theo thời gian.")
add_body("FuncAnimation gọi hàm vẽ lại biểu đồ mỗi 1000ms (1 giây), cho phép người dùng quan sát xu hướng kết nối mạng một cách trực quan mà không cần tải lại trang.")

add_section_title("2.6. Framework Flask và RESTful API")
add_body("Flask là micro web framework trong Python, nhẹ và linh hoạt, phù hợp để xây dựng REST API [7]. Đề tài sử dụng Flask để tạo 2 endpoint API trên Azure:")
add_bullet("POST /api/log: Nhận dữ liệu JSON từ ứng dụng Desktop, ghi vào file CSV trên Azure File Share.")
add_bullet("GET /api/csv: Trả về nội dung file CSV cho người dùng hoặc giảng viên xem trực tiếp trên trình duyệt.")
add_body("Cơ chế xác thực Bearer Token được áp dụng: mỗi request phải kèm header Authorization: Bearer <token> hoặc query parameter ?token=<token> để được chấp nhận.")

add_section_title("2.7. Điện toán đám mây với Microsoft Azure")
add_subsection_title("2.7.1. Azure App Service")
add_body("Azure App Service là dịch vụ PaaS (Platform as a Service) cho phép triển khai ứng dụng web mà không cần quản lý máy chủ vật lý [8]. Ứng dụng Flask được deploy trực tiếp bằng lệnh az webapp deployment source config-zip, Azure tự động nhận diện và chạy ứng dụng Python qua Gunicorn.")

add_subsection_title("2.7.2. Azure Storage Account và File Share")
add_body("Azure File Share là dịch vụ lưu trữ file trên đám mây theo giao thức SMB (Server Message Block). Bằng cách mount (gắn kết) File Share vào Web App, tất cả file được ghi vào thư mục mount sẽ lưu trực tiếp trên đám mây, đảm bảo dữ liệu tồn tại độc lập với vòng đời ứng dụng [8].")

add_subsection_title("2.7.3. Azure for Students")
add_body("Chương trình Azure for Students cung cấp $100 USD credit miễn phí cho sinh viên đại học có địa chỉ email trường. Credit được khấu trừ tự động vào chi phí sử dụng dịch vụ Azure, không yêu cầu thẻ tín dụng [9].")

add_page_break()

# ============================================================
# CHƯƠNG 3. KẾT QUẢ THỰC NGHIỆM
# ============================================================
add_chapter_title("Chương 3. KẾT QUẢ THỰC NGHIỆM")

add_section_title("3.1. Thiết kế hệ thống")
add_subsection_title("3.1.1. Kiến trúc tổng quan")
add_body("Hệ thống được thiết kế theo mô hình Client – Server, gồm hai thành phần chính:")
add_bullet("Client (Desktop App): monitor.py chạy trên máy tính người dùng, thực hiện ping, hiển thị GUI, vẽ biểu đồ và gửi dữ liệu giám sát lên Cloud.")
add_bullet("Server (Cloud API): azure_app/app.py chạy trên Azure App Service, tiếp nhận dữ liệu JSON và lưu trữ bền vững trên Azure File Share.")

# Architecture diagram as code block
arch = """[Máy tính người dùng]                    [Microsoft Azure Cloud]
+----------------------+                  +----------------------------+
|   monitor.py (GUI)   |  HTTP POST JSON  |  Azure App Service (B1)    |
|  - Ping (ping3)      |----------------->|  Flask API (app.py)        |
|  - LAN Scanner       |  Bearer Token    |  POST /api/log             |
|  - Live Graph        |                  |  GET  /api/csv             |
|  - CSV Backup        |                  |         | mount             |
|  log.csv (local)     |                  |  Azure File Share          |
+----------------------+                  |  csvreports/log.csv        |
                                          +----------------------------+"""
add_code_block(arch)
add_figure_caption("Hình 3.1. Kiến trúc tổng quan hệ thống giám sát mạng")

add_subsection_title("3.1.2. Cấu trúc thư mục dự án")
add_table(
    ["File / Thư mục", "Vai trò"],
    [
        ["monitor.py", "Ứng dụng Desktop – GUI giám sát mạng (Tkinter + Matplotlib)"],
        ["azure_app/app.py", "Flask API – Tiếp nhận và lưu trữ log trên Cloud"],
        ["azure_app/requirements.txt", "Thư viện Python cần thiết cho Cloud API"],
        ["log.csv", "File log dự phòng local (14.765+ bản ghi thu thập)"],
        ["README.md", "Tài liệu hướng dẫn dự án"],
    ],
    caption="Bảng 3.1. Mô tả các file trong dự án"
)

add_subsection_title("3.1.3. Luồng xử lý chính")
add_body("Khi người dùng nhấn nút Start:")
add_bullet("Ứng dụng tạo một thread nền (daemon=True) để ping liên tục mỗi 2 giây.")
add_bullet("Mỗi lần ping, kết quả được xử lý đồng thời: cập nhật GUI, lưu dữ liệu biểu đồ, ghi log.csv local, gửi lên Azure Cloud API.")
add_bullet("Trên Cloud, Flask nhận JSON và append vào log.csv trên Azure File Share.")

add_section_title("3.2. Cài đặt chi tiết – Ứng dụng Desktop (monitor.py)")

add_subsection_title("3.2.1. Giao diện người dùng (GUI)")
add_body("(Chèn ảnh chụp giao diện ứng dụng tại đây)")
add_figure_caption("Hình 3.2. Giao diện ứng dụng Network Monitor PRO")

add_table(
    ["Widget", "Chức năng"],
    [
        ["Entry (ô nhập IP)", "Nhập địa chỉ IP/hostname cần giám sát (mặc định: 8.8.8.8)"],
        ["Label Status", "Hiển thị Online (xanh) / Offline (đỏ) theo thời gian thực"],
        ["Label Time", "Hiển thị thời gian của lần kiểm tra gần nhất"],
        ["Button Start", "Bắt đầu giám sát, tự khóa khi đang chạy (chống bấm đúp)"],
        ["Button Stop", "Dừng giám sát, mở khóa lại nút Start"],
        ["Button Show Live Graph", "Mở cửa sổ biểu đồ Matplotlib thời gian thực"],
        ["Button Scan Quick LAN", "Quét nhanh 10 IP đầu tiên trong mạng LAN"],
        ["Text Area", "Hiển thị kết quả quét LAN"],
    ],
    caption="Bảng 3.2. Các thành phần giao diện GUI"
)

add_subsection_title("3.2.2. Chức năng Ping và giám sát liên tục")
code_ping = """def check_network():
    global running, times, status_values
    while running:
        ip = entry_ip.get()
        result = ping(ip, timeout=1)   # ping3 library - ICMP

        status = 1 if result else 0
        status_text = "Online" if result else "Offline"
        current_time = time.strftime("%H:%M:%S")

        # Thread-safe: cap nhat UI tu luong chinh
        root.after(0, update_gui, status_text, current_time)

        # Ghi log local du phong
        with open("log.csv", "a", newline="") as f:
            csv.writer(f).writerow([ip, status_text, current_time])

        # Gui len Cloud Azure (fail-safe)
        try:
            requests.post(AZURE_API_URL, json={
                "ip": ip, "status": status_text, "time": current_time
            }, headers={"Authorization": "Bearer Anna-Secret-Token-2026"}, timeout=2)
        except Exception as e:
            print(f"Loi ket noi Server Azure: {e}")

        time.sleep(2)"""
add_code_block(code_ping)

add_body("Điểm kỹ thuật quan trọng:")
add_bullet("root.after(0, ...): Đẩy cập nhật UI về luồng chính, tránh crash do Tkinter không thread-safe.")
add_bullet("daemon=True: Thread tự hủy khi cửa sổ ứng dụng đóng.")
add_bullet("Fail-safe: Exception khi gửi Cloud được bắt và bỏ qua – ứng dụng tiếp tục chạy bình thường.")

add_subsection_title("3.2.3. Biểu đồ thời gian thực")
add_body("(Chèn ảnh chụp biểu đồ Matplotlib tại đây)")
add_figure_caption("Hình 3.3. Biểu đồ giám sát mạng thời gian thực (FuncAnimation)")
add_code_block("ani = FuncAnimation(fig, animate, interval=1000, cache_frame_data=False)")
add_body("Biểu đồ cập nhật mỗi 1 giây, giới hạn 60 điểm dữ liệu gần nhất, trục Y chỉ 2 giá trị: Offline (0) và Online (1).")

add_subsection_title("3.2.4. Quét mạng LAN")
add_body("(Chèn ảnh chụp kết quả quét LAN tại đây)")
add_figure_caption("Hình 3.4. Kết quả quét nhanh mạng LAN")
code_lan = """def scan_lan_thread():
    base_ip = entry_ip.get().rsplit('.', 1)[0]
    for i in range(1, 11):
        ip = f"{base_ip}.{i}"
        result = ping(ip, timeout=0.5)
        if result:
            root.after(0, lambda cur=ip:
                result_text.insert(tk.END, f"[+] {cur} is ONLINE\\n"))"""
add_code_block(code_lan)
add_body("Chức năng quét LAN sử dụng vòng lặp ping qua 10 IP đầu tiên trong cùng subnet, chạy trong luồng nền để không làm đóng băng giao diện.")

add_section_title("3.3. Cài đặt chi tiết – API Cloud (azure_app/app.py)")

add_subsection_title("3.3.1. Endpoint POST /api/log – Nhận và lưu dữ liệu")
code_post = """@app.route("/api/log", methods=["POST"])
def log_data():
    if not check_auth(request):
        return jsonify({"error": "Unauthorized"}), 401
    data = request.json
    os.makedirs(FILE_SHARE_PATH, exist_ok=True)
    file_exists = os.path.isfile(CSV_FILE_PATH)
    with open(CSV_FILE_PATH, "a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        if not file_exists:
            writer.writerow(["ip", "status", "time"])
        writer.writerow([data.get("ip"), data.get("status"), data.get("time")])
    return jsonify({"status": "ok"})"""
add_code_block(code_post)

add_subsection_title("3.3.2. Endpoint GET /api/csv – Xem dữ liệu trên trình duyệt")
add_body("(Chèn ảnh chụp trình duyệt hiển thị CSV tại đây)")
add_figure_caption("Hình 3.5. Dữ liệu CSV trả về qua trình duyệt")
add_body("API hỗ trợ 2 kiểu xác thực linh hoạt:")
add_bullet("Header: Authorization: Bearer Anna-Secret-Token-2026")
add_bullet("Query param: ?token=Anna-Secret-Token-2026 (mở thẳng trên trình duyệt)")

add_section_title("3.4. Triển khai lên Azure")

add_subsection_title("3.4.1. Cấu hình tài nguyên Azure")
add_table(
    ["Tài nguyên", "Giá trị", "Lý do chọn"],
    [
        ["Resource Group", "RG_Network_Monitor", "Nhóm quản lý toàn bộ tài nguyên"],
        ["Location", "Malaysia West", "Gần Việt Nam, phù hợp tài khoản sinh viên"],
        ["Storage Account", "stloganna2026 (Standard_LRS)", "Lưu trữ CSV bền vững chi phí thấp"],
        ["File Share", "csvreports", "Giao thức SMB mount vào Web App"],
        ["App Service Plan", "ASP-NetworkMonitor (B1, Linux)", "Gói Basic – đủ cho API nhẹ"],
        ["Web App", "network-monitor-pro-anna", "Python 3.12 runtime"],
        ["Mount Path", "/mounts/csvreports", "Đường dẫn mount File Share trong app"],
    ],
    caption="Bảng 3.3. Cấu hình tài nguyên Azure đã triển khai"
)

add_subsection_title("3.4.2. Lưu trữ bền vững với Azure File Share")
add_body("(Chèn ảnh chụp Azure Portal → File Share csvreports tại đây)")
add_figure_caption("Hình 3.6. File Share csvreports trên Azure Portal")
add_body("Khi Flask ghi vào /mounts/csvreports/log.csv, dữ liệu được lưu trực tiếp trên Azure File Share, độc lập với vòng đời Web App. Redeploy hay restart không làm mất dữ liệu.")

add_section_title("3.5. Kết quả thu được")

add_subsection_title("3.5.1. Thống kê dữ liệu thu thập")
add_table(
    ["Thông số", "Giá trị"],
    [
        ["Cloud API URL", "https://network-monitor-pro-anna.azurewebsites.net"],
        ["Tổng bản ghi thu thập", "14.765 bản ghi"],
        ["Các IP đã giám sát", "8.8.8.8, 1.1.1.1, 192.168.1.90, 10.12.3.221, ..."],
        ["Chu kỳ ping", "2 giây/lần"],
        ["Gói Azure", "Azure for Students ($100 credit)"],
    ],
    caption="Bảng 3.4. Thống kê dữ liệu thực nghiệm"
)

add_subsection_title("3.5.2. Mẫu dữ liệu CSV")
add_body("Local log.csv (14.765+ bản ghi):")
add_code_block("""8.8.8.8,Online,02:05:00
8.8.8.8,Online,02:05:02
192.168.1.90,Online,04:43:10
1.1.1.1,Online,04:52:02
8.8.8.8,Offline,10:59:20   <- phat hien mat ket noi
8.8.8.8,Online,10:59:23    <- ket noi phuc hoi""")

add_body("Cloud log.csv (Azure File Share):")
add_code_block("""ip,status,time
8.8.8.8,Online,20:47:16
8.8.8.8,Online,20:47:18
8.8.8.8,Online,20:47:20""")

add_subsection_title("3.5.3. Đánh giá kết quả tổng thể")
add_table(
    ["Chỉ tiêu", "Kết quả", "Đánh giá"],
    [
        ["GUI hoạt động ổn định, không freeze", "✅ Đạt", "Thread nền + root.after()"],
        ["Biểu đồ real-time cập nhật đúng", "✅ Đạt", "FuncAnimation 1s/lần"],
        ["Quét LAN phát hiện thiết bị", "✅ Đạt", "Phát hiện đúng Online/Offline"],
        ["Ghi log local 14.765+ bản ghi", "✅ Đạt", "CSV append mode hoạt động"],
        ["API Cloud nhận và lưu dữ liệu", "✅ Đạt", "Azure File Share mount OK"],
        ["Bảo mật Bearer Token", "✅ Đạt", "401 khi không có token"],
        ["File Share bền vững qua redeploy", "✅ Đạt", "Dữ liệu không mất sau restart"],
        ["Fail-safe khi mất kết nối Cloud", "✅ Đạt", "App tiếp tục chạy bình thường"],
    ],
    caption="Bảng 3.5. Tổng kết đánh giá các chỉ tiêu"
)

add_page_break()

# ============================================================
# CHƯƠNG 4. KẾT LUẬN VÀ KIẾN NGHỊ
# ============================================================
add_chapter_title("Chương 4. KẾT LUẬN VÀ KIẾN NGHỊ")

add_section_title("4.1. Kết luận")
add_body("Đề tài đã hoàn thành xây dựng hệ thống giám sát mạng hoàn chỉnh gồm ứng dụng Desktop có GUI và API Cloud trên Azure.")
add_body("Các kết quả đạt được:")
add_bullet("Về giao diện người dùng: Ứng dụng Tkinter với đầy đủ chức năng: giám sát real-time, biểu đồ trực quan, quét LAN, cập nhật màu sắc trạng thái (xanh – Online, đỏ – Offline).")
add_bullet("Về kỹ thuật lập trình: Áp dụng đa luồng (threading) bảo đảm UI không freeze. Xử lý đúng thread-safe với root.after(). Fail-safe exception handling cho kết nối Cloud.")
add_bullet("Về tích hợp Cloud: REST API Flask trên Azure App Service nhận dữ liệu liên tục và lưu bền vững trên Azure File Share. Hệ thống đã thu thập 14.765+ bản ghi log thành công.")
add_bullet("Về bảo mật: Bearer Token Authentication ngăn truy cập trái phép. Hỗ trợ 2 kiểu xác thực (Header + Query Param) linh hoạt.")
add_bullet("Về chi phí: Toàn bộ chi phí vận hành nằm trong phạm vi Azure for Students ($100 credit), không phát sinh chi phí thực.")

add_section_title("4.2. Kiến nghị và hướng phát triển")
add_body("Hạn chế hiện tại:")
add_table(
    ["Hạn chế", "Nguyên nhân", "Mức độ ảnh hưởng"],
    [
        ["Chưa đóng gói .exe/.app", "Chưa dùng PyInstaller", "Trung bình – cần cài Python"],
        ["Chưa có cảnh báo realtime", "Chưa tích hợp Telegram/Slack", "Cao – ảnh hưởng thực tiễn"],
        ["Quét LAN giới hạn 10 IP", "Thiết kế ban đầu đơn giản", "Thấp"],
        ["Cloud CSV chưa có giao diện web", "Chỉ trả raw CSV", "Trung bình"],
    ],
    caption="Bảng 4.1. Hạn chế hiện tại"
)

add_body("Hướng phát triển trong tương lai:")
add_bullet("Alert System: Tích hợp Telegram Bot API hoặc Slack Webhook để gửi cảnh báo tức thời khi phát hiện host down.")
add_bullet("Dashboard Web: Xây dựng giao diện web bằng Flask + Chart.js hiển thị biểu đồ latency theo thời gian thực.")
add_bullet("Đóng gói ứng dụng: Sử dụng PyInstaller để đóng gói thành file .exe (Windows) hoặc .app (macOS).")
add_bullet("Mở rộng LAN Scanner: Quét toàn bộ /24 subnet (254 IP) với ThreadPoolExecutor.")
add_bullet("Database Integration: Thay CSV bằng SQLite hoặc Azure Cosmos DB để hỗ trợ truy vấn phức tạp.")

add_page_break()

# ============================================================
# TÀI LIỆU THAM KHẢO
# ============================================================
add_chapter_title("TÀI LIỆU THAM KHẢO")

refs = [
    "[1] Subramanian, M. (2000). Network Management: Principles and Practice. Addison-Wesley.",
    "[2] Python Software Foundation (2023). Python 3 Standard Library Documentation. https://docs.python.org/3/library/",
    "[3] Postel, J. (1981). Internet Control Message Protocol (RFC 792). IETF. https://www.rfc-editor.org/rfc/rfc792",
    "[4] Postel, J. (1981). Transmission Control Protocol (RFC 793). IETF. https://www.rfc-editor.org/rfc/rfc793",
    "[5] Mockapetris, P. (1987). Domain Names – Specification (RFC 1035). IETF. https://www.rfc-editor.org/rfc/rfc1035",
    "[6] Hunter, J.D. (2007). Matplotlib: A 2D Graphics Environment. Computing in Science & Engineering, 9(3), 90–95.",
    "[7] Grinberg, M. (2018). Flask Web Development. O'Reilly Media.",
    "[8] Microsoft Corporation (2024). Azure App Service Documentation. https://learn.microsoft.com/en-us/azure/app-service/",
    "[9] Microsoft Corporation (2024). Azure for Students. https://azure.microsoft.com/en-us/free/students/",
]
for r in refs:
    add_body(r)

add_page_break()

# ============================================================
# PHỤ LỤC
# ============================================================
add_chapter_title("PHỤ LỤC")

add_section_title("Phụ lục A. Mã nguồn monitor.py")
add_body("Xem file: monitor.py trong thư mục gốc dự án.")

# Read actual source
monitor_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "monitor.py")
if os.path.exists(monitor_path):
    with open(monitor_path, "r", encoding="utf-8") as f:
        src = f.read()
    add_code_block(src)

add_section_title("Phụ lục B. Mã nguồn azure_app/app.py")
add_body("Xem file: azure_app/app.py trong thư mục dự án.")

app_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "azure_app", "app.py")
if os.path.exists(app_path):
    with open(app_path, "r", encoding="utf-8") as f:
        src2 = f.read()
    add_code_block(src2)

add_section_title("Phụ lục C. Mẫu dữ liệu log.csv")
add_code_block("""ip,status,time
8.8.8.8,Online,02:05:00
8.8.8.8,Online,02:05:02
192.168.1.90,Online,04:43:10
1.1.1.1,Online,04:52:02
10.12.3.221,Online,10:03:06
8.8.8.8,Offline,10:59:20
8.8.8.8,Online,10:59:23""")

add_section_title("Phụ lục D. Lệnh Azure CLI triển khai")
add_code_block("""# Tao Resource Group
az group create --name RG_Network_Monitor --location malaysiawest

# Tao Storage + File Share
az storage account create --name stloganna2026 \\
  --resource-group RG_Network_Monitor --location malaysiawest --sku Standard_LRS
az storage share create --name csvreports --account-name stloganna2026

# Tao App Service + Web App
az appservice plan create --name ASP-NetworkMonitor \\
  --resource-group RG_Network_Monitor --location malaysiawest --sku B1 --is-linux
az webapp create --resource-group RG_Network_Monitor \\
  --plan ASP-NetworkMonitor --name network-monitor-pro-anna --runtime "PYTHON:3.12"

# Mount File Share
az webapp config storage-account add \\
  --resource-group RG_Network_Monitor --name network-monitor-pro-anna \\
  --custom-id csvmount --storage-type AzureFiles \\
  --share-name csvreports --account-name stloganna2026 \\
  --access-key <KEY> --mount-path /mounts/csvreports

# Deploy code
cd azure_app && zip -r ../deploy.zip .
az webapp deployment source config-zip \\
  --resource-group RG_Network_Monitor \\
  --name network-monitor-pro-anna --src deploy.zip""")

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "BaoCao_DoAn.docx")
doc.save(output_path)
print(f"✅ Đã tạo thành công: {output_path}")
